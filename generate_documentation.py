from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


def set_default_font(document: Document, font_name: str = "Calibri", font_size: int = 11) -> None:
    styles = document.styles
    normal_style = styles["Normal"]
    normal_style.font.name = font_name
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    normal_style.font.size = Pt(font_size)


def add_heading(document: Document, text: str, level: int = 1):
    heading = document.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Calibri"
    return heading


def main() -> None:
    out_path = Path(r"C:\Users\aboof\roadsos\documentation.docx")

    document = Document()
    set_default_font(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RoadSoS — AI Emergency Response System")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(24)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("National Road Safety Hackathon 2026 | IIT Madras (CoERS)")
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(12)

    document.add_paragraph("")

    add_heading(document, "Section 1: Project Overview", level=1)
    document.add_paragraph(
        "RoadSoS is an AI-powered emergency accident triage and response system designed to reduce emergency response times during road accidents. "
        "The system accepts voice or text descriptions of accidents, extracts key information using NLP, predicts severity using machine learning, "
        "and locates the nearest appropriate emergency resources using real OpenStreetMap data."
    )

    add_heading(document, "Section 2: Software Packages Used", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Package Name"
    header_cells[1].text = "Version"
    header_cells[2].text = "Purpose"
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    packages = [
        ("streamlit", "1.x", "Web application framework"),
        ("spacy", "3.x", "Natural Language Processing"),
        ("en_core_web_sm", "3.8.0", "spaCy English language model"),
        ("openai-whisper", "latest", "Speech to text transcription"),
        ("osmnx", "2.x", "OpenStreetMap data fetching"),
        ("folium", "0.20.0", "Interactive map generation"),
        ("scikit-learn", "1.8.0", "Machine learning (Random Forest)"),
        ("pandas", "3.x", "Data manipulation"),
        ("numpy", "latest", "Numerical computation"),
        ("geopy", "2.x", "Geocoding and distance calculation"),
        ("networkx", "3.x", "Graph-based routing"),
        ("thefuzz", "latest", "Fuzzy string matching for NLP"),
        ("google-generativeai", "0.8.6", "Gemini Vision API for image severity"),
        ("python-dotenv", "1.2.2", "Environment variable management"),
        ("sounddevice", "latest", "Audio recording"),
        ("scipy", "latest", "Audio file processing"),
        ("matplotlib", "latest", "Data visualization"),
        ("seaborn", "latest", "Statistical visualization"),
        ("xgboost", "latest", "Gradient boosting (explored)"),
        ("folium.plugins HeatMap", "included", "Blackspot heatmap visualization"),
    ]

    for package_name, version, purpose in packages:
        row_cells = table.add_row().cells
        row_cells[0].text = package_name
        row_cells[1].text = version
        row_cells[2].text = purpose

    add_heading(document, "Section 3: System Architecture", level=1)
    architecture_items = [
        ("Input Layer", "Voice (Whisper) + Text + Image (Gemini Vision)"),
        ("Intelligence Layer", "spaCy NLP + Fuzzy matching + Random Forest ML"),
        ("Resource Layer", "OSMnx + Geopy for real-time hospital/police/workshop location"),
        ("Output Layer", "Streamlit UI + Folium maps + Emergency contact links"),
    ]
    for title_text, description in architecture_items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(f"{title_text} — {description}")

    add_heading(document, "Section 4: Assumptions Made", level=1)
    assumptions = [
        "Hospital and police station data is sourced from OpenStreetMap and may not include all private facilities",
        "Accident severity is predicted based on keyword analysis of user-provided text descriptions",
        "Location detection is limited to Chennai, Tamil Nadu area names",
        "Blackspot map uses simulated accident data based on known high-risk zones in Chennai",
        "Voice transcription accuracy depends on microphone quality and background noise",
        "Image severity analysis requires internet connection for Gemini API calls",
        "The ML model was trained on state-level MORTH road accident data (2022) and is used for demonstration purposes",
        "Emergency contact numbers shown (108, 100, 101, 104) are standard Indian national helpline numbers",
    ]
    for assumption in assumptions:
        document.add_paragraph(assumption, style="List Bullet")

    add_heading(document, "Section 5: Data Sources", level=1)
    data_sources = [
        ("Road accident data", "data.gov.in — State/UT wise Road Accidents 2022"),
        ("Map data", "OpenStreetMap contributors via OSMnx"),
        ("Image analysis", "Google Gemini 1.5 Flash API"),
    ]
    for label, source in data_sources:
        paragraph = document.add_paragraph()
        bold_run = paragraph.add_run(f"{label}: ")
        bold_run.bold = True
        paragraph.add_run(source)

    document.save(out_path)
    print(f"Created {out_path}")


if __name__ == "__main__":
    main()